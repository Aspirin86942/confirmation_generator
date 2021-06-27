import xlrd
import re
from docx import Document
import os
import shutil
from comtypes.client import CreateObject


# 读取excel
class ReadExecl(object):

    def __init__(self, input_path, sheet_name):
        self.input_path = input_path
        self.sheet_name = sheet_name

    def read_excel(self):
        # 打开excel表，填写路径
        book = xlrd.open_workbook(self.input_path)
        # 找到sheet页
        table = book.sheet_by_name(self.sheet_name)
        # 获取总行数总列数
        row_num = table.nrows
        col_num = table.ncols
        ls = []
        key = table.row_values(0)
        # 这是第一行数据，作为字典的key值
        if row_num <= 1:
            # print("没数据")
            return ls
        else:
            j = 1
            for i in range(row_num - 1):
                d = {}
                values = table.row_values(j)
                for x in range(col_num):
                    # 把key值对应的value赋值给key，每行循环
                    d[key[x]] = values[x]
                j += 1
                # 把字典加到列表中
                ls.append(d)
            return ls

    def distinct_excel(self, dict):
        main_key = []
        catalogue_dict = {}
        for i in range(len(dict)):
            main_key.append(dict[i]["索引号"])
        for item in main_key:
            catalogue_dict.update({item: ""})
        keys = []
        for key in catalogue_dict:
            keys.append(key)
        return keys

    def make_table(self, key, dict):
        table = []
        for j in range(len(dict)):
            if dict[j]["索引号"] == key:
                table.append(dict[j])
        return table


# 生成word
class MakeWord(object):

    def __init__(self, output_path, table_bank_deposit, table_bank_loan, table_account_cancel, table_entrusted_loan):
        self.output_path = output_path
        self.table_bank_deposit = table_bank_deposit
        self.table_bank_loan = table_bank_loan
        self.table_account_cancel = table_account_cancel
        self.table_entrusted_loan = table_entrusted_loan

    # 思路是同一个索引号的函证在一个函数中完成，蠢是蠢了点，但是我觉得不错
    def make_doc(self):
        global replace_dict  # 设定基本信息字典
        replace_dict = {}
        document = Document('银行询证函(模板) - 开除函证中心.docx')

        # 设定表
        bank_deposit = document.tables[0]
        bank_loan = document.tables[1]
        account_cancel = document.tables[2]
        entrusted_loan = document.tables[3]

        # 判断是否有银行存款,并对银行存款进行填充
        MakeWord.filling_bank_deposit(self, bank_deposit=bank_deposit, document=document)

        # 判断是否有银行借款，并对银行借款进行填充
        MakeWord.filling_bank_loan(self, bank_loan=bank_loan, document=document)

        # 判断是否有银行销户
        MakeWord.filling_account_cancel(self, account_cancel=account_cancel, document=document)

        # 判断是否有本公司作为贷款方的委托贷款
        MakeWord.filling_entrusted_loan(self, entrusted_loan=entrusted_loan, document=document)

        # 保存函证
        document.save(r'.\word版函证\%s.docx' % self.output_path)

        # 进度提示
        print("%s生成完毕" % self.output_path)

    def filling_bank_deposit(self, bank_deposit, document):
        # 判断是否有银行存款
        if self.table_bank_deposit:

            # 比较一下所填内容和word中银行存款的表格大小
            if len(bank_deposit.rows) - 1 <= len(self.table_bank_deposit):
                # 增加行
                for j in range(len(self.table_bank_deposit) - len(bank_deposit.rows) + 1):
                    bank_deposit.add_row()
            # 对doc进行银行存款数据填充
            for i in range(1, len(bank_deposit.rows)):
                bank_deposit.cell(i, 0).text = self.table_bank_deposit[i - 1]["账户（公司）名称"]
                bank_deposit.cell(i, 1).text = self.table_bank_deposit[i - 1]["银行账号"]
                bank_deposit.cell(i, 2).text = self.table_bank_deposit[i - 1]["币种"]
                bank_deposit.cell(i, 3).text = Assemble.reg_percent(self.table_bank_deposit[i - 1]["利率"])
                bank_deposit.cell(i, 4).text = self.table_bank_deposit[i - 1]["账户类型"]
                bank_deposit.cell(i, 5).text = Assemble.reg_balance(self.table_bank_deposit[i - 1]["余额"])
                bank_deposit.cell(i, 6).text = self.table_bank_deposit[i - 1]["是否属于资金归集（资金池或其他资金管理）账户"]
                bank_deposit.cell(i, 7).text = Assemble.reg_date(self.table_bank_deposit[i - 1]["起止日期"])
                bank_deposit.cell(i, 8).text = Assemble.reg_date(self.table_bank_deposit[i - 1]["终止日期"])
                bank_deposit.cell(i, 9).text = self.table_bank_deposit[i - 1]["是否用于担保或存在其他使用限制"]
                bank_deposit.cell(i, 10).text = self.table_bank_deposit[i - 1]["备注"]
            # 函证基本信息填充部分
            replace_dict.update({
                "数据1": self.table_bank_deposit[0]["索引号"],
                "数据2": self.table_bank_deposit[0]["银行名称"],
                "数据3": self.table_bank_deposit[0]["银行账号"],
                "数据4": self.table_bank_deposit[0]["账户（公司）名称"],
                "数据5": self.table_bank_deposit[0]["年度"]
            })
            # 开始替换
            Assemble.hz_replace(document=document, replace_dict=replace_dict)

    def filling_bank_loan(self, bank_loan, document):
        # 判断是否有银行借款
        if self.table_bank_loan:
            # 比较一下所填内容和word中银行借款的表格大小
            if len(bank_loan.rows) - 1 <= len(self.table_bank_loan):
                # 增加行
                for L in range(len(self.table_bank_loan) - len(bank_loan.rows) + 1):
                    bank_loan.add_row()
            # 对doc进行银行借款数据填充
            for k in range(1, len(bank_loan.rows)):
                bank_loan.cell(k, 0).text = self.table_bank_loan[k - 1]["借款人名称"]
                bank_loan.cell(k, 1).text = self.table_bank_loan[k - 1]["银行账号"]
                bank_loan.cell(k, 2).text = self.table_bank_loan[k - 1]["币种"]
                bank_loan.cell(k, 3).text = Assemble.reg_balance(self.table_bank_loan[k - 1]["余额"])
                bank_loan.cell(k, 4).text = Assemble.reg_date(self.table_bank_loan[k - 1]["借款日期"])
                bank_loan.cell(k, 5).text = Assemble.reg_date(self.table_bank_loan[k - 1]["到期日期"])
                bank_loan.cell(k, 6).text = Assemble.reg_percent(self.table_bank_loan[k - 1]["利率"])
                bank_loan.cell(k, 7).text = self.table_bank_loan[k - 1]["抵（质）押品/担保人"]
                bank_loan.cell(k, 8).text = self.table_bank_loan[k - 1]["备注"]

            if not bool(replace_dict):
                # 函证基本信息填充部分
                replace_dict.update({
                    "数据1": self.table_bank_loan[0]["索引号"],
                    "数据2": self.table_bank_loan[0]["银行名称"],
                    "（账号：数据3）": "",
                    "数据4": self.table_bank_loan[0]["借款人名称"],
                    "数据5": self.table_bank_loan[0]["年度"]
                })
                # 开始替换
                Assemble.hz_replace(document=document, replace_dict=replace_dict)

    def filling_account_cancel(self, account_cancel, document):
        if self.table_account_cancel:
            # 比较一下所填内容和word中银行销户的表格大小
            if len(account_cancel.rows) - 1 <= len(self.table_account_cancel):
                # 增加行
                for L in range(len(self.table_account_cancel) - len(account_cancel.rows) + 1):
                    account_cancel.add_row()
            # 对doc进行银行销户数据填充
            for k in range(1, len(account_cancel.rows)):
                account_cancel.cell(k, 0).text = self.table_account_cancel[k - 1]["账户名称"]
                account_cancel.cell(k, 1).text = self.table_account_cancel[k - 1]["银行账号"]
                account_cancel.cell(k, 2).text = self.table_account_cancel[k - 1]["币种"]
                account_cancel.cell(k, 3).text = Assemble.reg_date(self.table_account_cancel[k - 1]["注销账户日"])

            if not bool(replace_dict):
                # 函证基本信息填充部分
                replace_dict.update({
                    "数据1": self.table_account_cancel[0]["索引号"],
                    "数据2": self.table_account_cancel[0]["银行名称"],
                    "（账号：数据3）": "",
                    "数据4": self.table_account_cancel[0]["账户名称"],
                    "数据5": self.table_account_cancel[0]["年度"]
                })
                # 开始替换
                Assemble.hz_replace(document=document, replace_dict=replace_dict)

    def filling_entrusted_loan(self, entrusted_loan, document):
        if self.table_entrusted_loan:
            # 比较一下所填内容和word中本公司作为贷款方的委托贷款的表格大小
            if len(entrusted_loan.rows) - 1 <= len(self.table_entrusted_loan):
                # 增加行
                for L in range(len(self.table_entrusted_loan) - len(entrusted_loan.rows) + 1):
                    entrusted_loan.add_row()
            # 对doc进行本公司作为贷款方的委托贷款数据填充
            for k in range(1, len(entrusted_loan.rows)):
                entrusted_loan.cell(k, 0).text = self.table_entrusted_loan[k - 1]["账户（公司）名称"]
                entrusted_loan.cell(k, 1).text = self.table_entrusted_loan[k - 1]["银行结算账号"]
                entrusted_loan.cell(k, 2).text = self.table_entrusted_loan[k - 1]["资金借入方"]
                entrusted_loan.cell(k, 3).text = self.table_entrusted_loan[k - 1]["币种"]
                entrusted_loan.cell(k, 4).text = Assemble.reg_percent(self.table_entrusted_loan[k - 1]["利率"] * 100)
                entrusted_loan.cell(k, 5).text = Assemble.reg_balance(self.table_entrusted_loan[k - 1]["余额"])
                entrusted_loan.cell(k, 6).text = self.table_entrusted_loan[k - 1]["贷款起止日期"]
                entrusted_loan.cell(k, 7).text = self.table_entrusted_loan[k - 1]["备注"]

            if not bool(replace_dict):
                # 函证基本信息填充部分
                replace_dict.update({
                    "数据1": self.table_entrusted_loan[0]["索引号"],
                    "数据2": self.table_entrusted_loan[0]["银行名称"],
                    "（账号：数据3）": "",
                    "数据4": self.table_entrusted_loan[0]["账户（公司）名称"],
                    "数据5": self.table_entrusted_loan[0]["年度"],
                })
                # print(replace_dict)
                # 开始替换
                Assemble.hz_replace(document=document, replace_dict=replace_dict)


# word转pdf
class Word2Pdf(object):
    def __init__(self):
        # word文档转化为pdf文档时使用的格式为17
        self.wdFormatPDF = 17
        self.wdToPDF = CreateObject("Word.Application")

    def wd_to_pdf(self, input_path, output_path):
        # 获取指定目录下面的所有文件
        files = os.listdir(input_path)
        # files = list_nohidden(folder)
        # print(files)
        # 获取word类型的文件放到一个列表里面
        wdfiles = [f for f in files if f.endswith((".doc", ".docx"))]
        # 去除word生成的隐藏文件
        wdfiles2 = [f for f in wdfiles if not f.startswith('~')]
        for wdfile in wdfiles2:
            # 将word文件放到指定的路径下面
            wdPath = os.path.join(input_path, wdfile)
            # print(wdfile)
            # 设置将要存放pdf文件的路径
            pdfPath = output_path + wdfile.split(".")[0] + '.pdf'
            # 判断是否已经存在对应的pdf文件，如果不存在就加入到存放pdf的路径内
            if pdfPath[-3:] != 'pdf':
                pdfPath = pdfPath + ".pdf"
            # 将word文档转化为pdf文件，先打开word所在路径文件，然后在处理后保存pdf文件，最后关闭
            pdfCreate = self.wdToPDF.Documents.Open(wdPath)
            pdfCreate.SaveAs(pdfPath, self.wdFormatPDF)
            pdfCreate.Close()
            print("正在生成%s" % pdfPath)


# 静态方法的集合
class Assemble:
    @staticmethod
    # 替换函证信息的方法
    def hz_replace(document, replace_dict):
        for para in document.paragraphs:
            for i in range(len(para.runs)):
                for key, value in replace_dict.items():
                    if key in para.runs[i].text:
                        para.runs[i].text = para.runs[i].text.replace(key, value)

    @staticmethod
    # 对余额数据进行规范
    def reg_balance(balance):
        a = '%.2f' % (float(balance))
        balance_format = re.sub(r"(\d)(?=(\d\d\d)+(?!\d))", r"\1,", a)
        return str(balance_format)

    @staticmethod
    # 对日期数据进行规范
    def reg_date(date):
        try:
            b = xlrd.xldate.xldate_as_datetime(date, 0).strftime('%Y-%m-%d')
            return str(b)
        except TypeError:
            b = date
            return str(b)

    @staticmethod
    # 对百分数数据进行规范
    def reg_percent(percent):
        if percent == "活期":
            return percent
        else:
            try:
                percent_f = float(percent*100)
                c = str('{:.4f}%'.format(percent_f))
            except ValueError:
                print("利率须为填写为浮点、活期或不填，否则自动填写为活期")
                c = "活期"
            return str(c)

    @staticmethod
    # 创建word文件夹
    def setup_doc_file():
        is_exists = os.path.exists(r".\word版函证")
        if not is_exists:
            os.makedirs(r".\word版函证")
        else:
            shutil.rmtree(r".\word版函证")
            os.makedirs(r".\word版函证")

    @staticmethod
    # 创建pdf文件夹
    def setup_pdf_file():
        is_exists = os.path.exists(r".\pdf版函证")
        if not is_exists:
            os.makedirs(r".\pdf版函证")
        else:
            shutil.rmtree(r".\pdf版函证")
            os.makedirs(r".\pdf版函证")


# 主程序
def main():
    # 设定工作目录
    work_path = os.getcwd()
    os.chdir(work_path)

    # 创建文件夹
    Assemble.setup_doc_file()
    Assemble.setup_pdf_file()

    # 对银行存款进行读取
    ReadExecl_bank_deposit = ReadExecl(input_path="input.xlsx", sheet_name="银行存款")
    dict_bank_deposit = ReadExecl_bank_deposit.read_excel()
    # 对银行存款其中不重复的索引号生成目录
    catalogue_bank_deposit = ReadExecl_bank_deposit.distinct_excel(dict_bank_deposit)

    # 对银行借款进行读取
    ReadExecl_bank_loan = ReadExecl(input_path="input.xlsx", sheet_name="银行借款")
    dict_bank_loan = ReadExecl_bank_loan.read_excel()
    # 对银行借款其中不重复的索引号生成目录
    catalogue_bank_loan = ReadExecl_bank_loan.distinct_excel(dict_bank_loan)

    # 对银行销户进行读取
    ReadExecl_account_cancel = ReadExecl(input_path="input.xlsx", sheet_name="银行销户")
    dict_account_cancel = ReadExecl_account_cancel.read_excel()
    # 对银行借款其中不重复的索引号生成目录
    catalogue_account_cancel = ReadExecl_account_cancel.distinct_excel(dict_account_cancel)

    # 对本公司作为贷款方的委托贷款进行读取
    ReadExecl_entrusted_loan = ReadExecl(input_path="input.xlsx", sheet_name="本公司作为贷款方的委托贷款")
    dict_entrusted_loan = ReadExecl_entrusted_loan.read_excel()
    # 对本公司作为贷款方的委托贷款其中不重复的索引号生成目录
    catalogue_entrusted_loan = ReadExecl_entrusted_loan.distinct_excel(dict_entrusted_loan)

    catalogue_all = list(set(catalogue_bank_deposit + catalogue_bank_loan + catalogue_account_cancel +
                             catalogue_entrusted_loan))
    catalogue_all.sort(reverse=False)
    for item in catalogue_all:
        # 生成单个函证数据文件
        single_bank_deposit = ReadExecl_bank_deposit.make_table(item, dict_bank_deposit)
        single_bank_loan = ReadExecl_bank_loan.make_table(item, dict_bank_loan)
        single_account_cancel = ReadExecl_account_cancel.make_table(item, dict_account_cancel)
        single_entrusted_loan = ReadExecl_entrusted_loan.make_table(item, dict_entrusted_loan)

        # 实例化MakeWord类
        MakeWord_1 = MakeWord(output_path=item, table_bank_deposit=single_bank_deposit, table_bank_loan=single_bank_loan
                              , table_account_cancel=single_account_cancel, table_entrusted_loan=single_entrusted_loan)
        MakeWord_1.make_doc()

    # 生成函证后自动生成pdf
    print("正在生成pdf")
    converter = Word2Pdf()
    input_path = work_path + r"\word版函证"
    output_path = work_path + r"\pdf版函证\\"
    converter.wd_to_pdf(input_path=input_path, output_path=output_path)
    print("pdf生成完毕")
    # 保证窗口在程序结束后不消失
    input('Press Enter to exit...')


if __name__ == '__main__':
    main()
